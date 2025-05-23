"""
文件处理服务 - 从PDF和Word文档中提取文本
File processor service for extracting text from PDF and Word documents
"""

import logging
import io
from typing import Optional

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    from docx import Document
except ImportError:
    Document = None

logger = logging.getLogger(__name__)


class FileProcessorService:
    """
    文件处理服务类，支持PDF和Word文档文本提取
    File processor service class supporting PDF and Word document text extraction
    """
    
    async def extract_text(self, file_content: bytes, filename: str) -> str:
        """
        从文件内容中提取文本
        Extract text from file content
        
        Args:
            file_content: 文件二进制内容 / File binary content
            filename: 文件名 / Filename
            
        Returns:
            str: 提取的文本内容 / Extracted text content
        """
        try:
            file_extension = filename.lower().split('.')[-1]
            
            if file_extension == 'pdf':
                return await self._extract_from_pdf(file_content)
            elif file_extension in ['doc', 'docx']:
                return await self._extract_from_word(file_content)
            else:
                # 尝试作为纯文本处理
                # Try to process as plain text
                try:
                    return file_content.decode('utf-8')
                except UnicodeDecodeError:
                    return file_content.decode('utf-8', errors='ignore')
                    
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {str(e)}")
            return f"[文件解析错误 / File parsing error: {str(e)}]"
    
    async def _extract_from_pdf(self, file_content: bytes) -> str:
        """
        从PDF文件中提取文本
        Extract text from PDF file
        """
        if pypdf is None:
            return "[PDF解析器未安装 / PDF parser not installed]"
        
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = pypdf.PdfReader(pdf_file)
            
            text_content = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(f"--- 第{page_num + 1}页 / Page {page_num + 1} ---")
                        text_content.append(page_text)
                except Exception as e:
                    logger.warning(f"Error extracting page {page_num + 1}: {str(e)}")
                    text_content.append(f"[第{page_num + 1}页解析失败 / Page {page_num + 1} parsing failed]")
            
            return "\n\n".join(text_content) if text_content else "[PDF内容为空 / PDF content is empty]"
            
        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}")
            return f"[PDF解析错误 / PDF parsing error: {str(e)}]"
    
    async def _extract_from_word(self, file_content: bytes) -> str:
        """
        从Word文档中提取文本
        Extract text from Word document
        """
        if Document is None:
            return "[Word解析器未安装 / Word parser not installed]"
        
        try:
            doc_file = io.BytesIO(file_content)
            doc = Document(doc_file)
            
            text_content = []
            
            # 提取段落文本
            # Extract paragraph text
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # 提取表格文本
            # Extract table text
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    text_content.append("--- 表格 / Table ---")
                    text_content.extend(table_text)
            
            return "\n\n".join(text_content) if text_content else "[Word文档内容为空 / Word document content is empty]"
            
        except Exception as e:
            logger.error(f"Error processing Word document: {str(e)}")
            return f"[Word文档解析错误 / Word document parsing error: {str(e)}]"
    
    def get_supported_extensions(self) -> list:
        """
        获取支持的文件扩展名
        Get supported file extensions
        """
        supported = ['txt']
        
        if pypdf is not None:
            supported.append('pdf')
        
        if Document is not None:
            supported.extend(['doc', 'docx'])
        
        return supported
    
    def is_supported_file(self, filename: str) -> bool:
        """
        检查文件是否支持
        Check if file is supported
        """
        file_extension = filename.lower().split('.')[-1]
        return file_extension in self.get_supported_extensions() 